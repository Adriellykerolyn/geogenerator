from flask import Flask, request, send_file, jsonify
import geopandas as gpd
import pandas as pd
from docx import Document
from flask_cors import CORS # Eu importo a lib para o JS poder falar com o Python
import io

# EU CRIO O APP PRIMEIRO
app = Flask(__name__)
# AGORA EU LIBERO O ACESSO
CORS(app) 

@app.route('/processar', methods=['POST'])
def processar_geodata():
    try:
        # 1. Eu recebo o arquivo e os dados do seu formulário
        file = request.files['arquivo']
        titulo = request.form.get('titulo', 'Relatório Geográfico')
        layout = request.form.get('layout', 'lista')

        # 2. Eu uso o GeoPandas para ler os dados geográficos reais
        gdf = gpd.read_file(file)
        
        # Eu removo a geometria (mapa) para criar a tabela de texto no Word
        df_dados = pd.DataFrame(gdf.drop(columns='geometry') if 'geometry' in gdf.columns else gdf)

        # 3. Eu crio o documento Word profissional
        doc = Document()
        doc.add_heading(titulo, 0)

        if layout == 'tabela':
            # Eu crio a tabela com o cabeçalho automático baseado nas colunas do arquivo
            table = doc.add_table(rows=1, cols=len(df_dados.columns))
            table.style = 'Table Grid'
            for i, column in enumerate(df_dados.columns):
                table.rows[0].cells[i].text = str(column)

            # Eu preencho com as informações das linhas do arquivo
            for _, row in df_dados.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
        else:
            # Layout de lista para relatórios descritivos
            for _, row in df_dados.iterrows():
                p = doc.add_paragraph()
                for col, val in row.items():
                    p.add_run(f"{col}: ").bold = True
                    p.add_run(f"{val}\n")

        # 4. Eu preparo o download para o seu navegador
        target = io.BytesIO()
        doc.save(target)
        target.seek(0)

        return send_file(target, as_attachment=True, download_name=f"{titulo}.docx")

    except Exception as e:
        return jsonify({"erro": str(e)}), 400

if __name__ == '__main__':
    app.run(debug=True)